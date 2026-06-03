#!/usr/bin/env python3
"""文档渲染函数 — 将结构化数据渲染为 .docx 元素。"""
import os, re
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from constants import (PAGE_TOP_MM, PAGE_BOTTOM_MM, PAGE_LEFT_MM, PAGE_RIGHT_MM,
                         PAGE_CONTENT_WIDTH_MM, SIZE_2, SIZE_3, SIZE_4)
from fonts import (resolve_font, set_cn_font, set_line_spacing_28,
                    set_first_line_indent, add_spacer, warn)
from hierarchy import classify_paragraph, STYLE_SALUTATION

_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_HEADING_SPLIT_RE = re.compile(
    r"^([一二三四五六七八九十]+、|[（(][一二三四五六七八九十]+[）)])"
)


def setup_page(doc: Document):
    """A4 页面设置。"""
    s = doc.sections[0]
    s.page_width = Mm(210)
    s.page_height = Mm(297)
    s.top_margin = Mm(PAGE_TOP_MM)
    s.bottom_margin = Mm(PAGE_BOTTOM_MM)
    s.left_margin = Mm(PAGE_LEFT_MM)
    s.right_margin = Mm(PAGE_RIGHT_MM)


def _add_page_text_run(para, text, font_name="宋体", size=SIZE_4):
    """添加页码文本 run，设置中文字体。"""
    run = para.add_run(text)
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.insert(0, rFonts)
    return run


def add_page_number(doc: Document):
    """页脚居中页码：— 1 — 宋体四号。"""
    s = doc.sections[0]
    footer = s.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    _add_page_text_run(para, "— ")
    run = _add_page_text_run(para, "")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    run._element.append(fc1)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    run._element.append(it)
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    run._element.append(fc2)
    _add_page_text_run(para, " —")


def add_title(doc: Document, text, available_fonts, warnings):
    """标题：方正小标宋简体 二号 居中。"""
    af = resolve_font("方正小标宋简体", available_fonts, warnings)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    set_line_spacing_28(para)
    para.paragraph_format.first_line_indent = Pt(0)
    run = para.add_run(text)
    set_cn_font(run, af, SIZE_2, False)


def add_body_paragraph(doc: Document, text, style, available_fonts, warnings,
                       alignment=None, no_indent=False, runs=None):
    """添加一段正文/标题段落。"""
    af = resolve_font(style["font_name"], available_fonts, warnings)
    bf = resolve_font("仿宋_GB2312", available_fonts, warnings)
    is_heading = style.get("is_heading", False)
    bold = style.get("bold", False)
    size = style.get("size", SIZE_3)
    para = doc.add_paragraph()
    para.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    set_line_spacing_28(para)
    if no_indent:
        para.paragraph_format.first_line_indent = Pt(0)
    else:
        set_first_line_indent(para)
    stripped = text.strip()
    if re.match(r"^\d+\.", stripped) and not re.match(r"^\d+\.\s", stripped):
        m = re.match(r"^(\d+\.)(.*)", stripped)
        if m:
            stripped = m.group(1) + " " + m.group(2)
    if runs:
        for r in runs:
            rt = r['text']
            rb = r.get('bold', False)
            run = para.add_run(rt)
            set_cn_font(run, af if is_heading else bf, size, bold or rb)
        return
    heading_split = None
    if is_heading:
        m = _HEADING_SPLIT_RE.match(stripped)
        if m:
            after_marker = stripped[m.end():]
            sm = re.search(r"[：。]", after_marker)
            if sm and len(after_marker[:sm.start()].strip()) >= 2:
                heading_split = m.end() + sm.end()
    if heading_split:
        prefix = stripped[:heading_split]
        suffix = stripped[heading_split:].lstrip()
        run1 = para.add_run(prefix)
        set_cn_font(run1, af, size, bold)
        if suffix:
            run2 = para.add_run(suffix)
            set_cn_font(run2, bf, size, False)
    else:
        run = para.add_run(stripped)
        set_cn_font(run, af, size, bold)


def _render_title(doc, data, available_fonts, warnings, body_font):
    title = data.get("标题", "").strip()
    if title:
        add_title(doc, title, available_fonts, warnings)


def _render_salutation(doc, data, available_fonts, warnings, body_font):
    sal = data.get("称谓", "").strip()
    if sal:
        add_body_paragraph(doc, sal, STYLE_SALUTATION, available_fonts, warnings, no_indent=True)


def _render_body(doc, data, available_fonts, warnings, body_font):
    for item in data.get("正文", []):
        if isinstance(item, dict) and "image" in item:
            img_path = item["image"]
            if not os.path.exists(img_path):
                warn(warnings, f"图片不存在: {img_path}")
                continue
            wm = min(item.get("width_mm", PAGE_CONTENT_WIDTH_MM), PAGE_CONTENT_WIDTH_MM)
            hm = item.get("height_mm", wm * 0.75)
            if item.get("width_mm") and item["width_mm"] > PAGE_CONTENT_WIDTH_MM:
                ratio = PAGE_CONTENT_WIDTH_MM / item["width_mm"]
                hm = item["height_mm"] * ratio
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Pt(0)
            run = para.add_run()
            try:
                run.add_picture(img_path, width=Mm(wm), height=Mm(hm))
            except Exception as e:
                warn(warnings, f"图片插入失败 ({os.path.basename(img_path)}): {e}")
                p = para._element
                p.getparent().remove(p)
            continue

        if isinstance(item, dict) and "table_xml" in item:
            try:
                clean_xml = re.sub(
                    r'\s+xmlns:wpsCustomData="http://www\.wps\.cn/officeDocument/2013/wpsCustomData"',
                    '', item["table_xml"]
                )
                tbl_elem = etree.fromstring(clean_xml)
                # 插入到 sectPr 之前（append 会放到 sectPr 之后导致表格跑末尾）
                sectPr = doc.element.body.find('{%s}sectPr' % _W_NS)
                if sectPr is not None:
                    sectPr.addprevious(deepcopy(tbl_elem))
                else:
                    doc.element.body.append(deepcopy(tbl_elem))
            except Exception as e:
                warn(warnings, f"表格 XML 注入失败: {e}")
            add_spacer(doc, "仿宋_GB2312", available_fonts, warnings, count=1)
            continue

        if isinstance(item, dict) and "table" in item:
            rows = item["table"]
            if not rows:
                continue
            max_cols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=max_cols)
            tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tbl.autofit = True
            tbl.style = "Table Grid"
            hfn = resolve_font("黑体", available_fonts, warnings)
            cfn = resolve_font("仿宋_GB2312", available_fonts, warnings)
            title_row = (len(rows) > 1 and len(rows[0]) == 1 and max_cols > 1)
            header_row_idx = 1 if title_row else 0
            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    cell = tbl.rows[ri].cells[ci]
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(cell_text)
                    if ri <= header_row_idx:
                        set_cn_font(run, hfn, Pt(10.5), False)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        set_cn_font(run, cfn, Pt(10.5), False)
                if title_row and ri == 0 and max_cols > 1:
                    tbl.rows[0].cells[0].merge(tbl.rows[0].cells[max_cols - 1])
            add_spacer(doc, "仿宋_GB2312", available_fonts, warnings, count=1)
            continue

        item_runs = None
        if isinstance(item, str):
            text = item.strip()
            pa = None
            ni = False
        else:
            text = item.get("text", "").strip()
            am = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                  "right": WD_ALIGN_PARAGRAPH.RIGHT,
                  "center": WD_ALIGN_PARAGRAPH.CENTER}
            pa = am.get(item.get("align"), None)
            ni = item.get("indent") is False
            item_runs = item.get("runs")
        if not text:
            continue
        style = classify_paragraph(text)
        add_body_paragraph(doc, text, style, available_fonts, warnings,
                          alignment=pa, no_indent=ni, runs=item_runs)


def _render_attachments(doc, data, available_fonts, warnings, body_font):
    atts = data.get("附件", [])
    if not atts:
        return
    add_spacer(doc, "仿宋_GB2312", available_fonts, warnings, count=1)
    for i, item in enumerate(atts, 1):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        set_line_spacing_28(para)
        set_first_line_indent(para)
        text = f"附件：{i}. {item}" if i == 1 else f"      {i}. {item}"
        run = para.add_run(text)
        set_cn_font(run, body_font, SIZE_3, False)


def _render_signature(doc, data, available_fonts, warnings, body_font):
    sig = data.get("署名", "").strip()
    dt = data.get("日期", "").strip()
    if not (sig or dt):
        return
    add_spacer(doc, "仿宋_GB2312", available_fonts, warnings, count=2)
    if sig:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        set_line_spacing_28(para)
        para.paragraph_format.first_line_indent = Pt(0)
        run = para.add_run(sig)
        set_cn_font(run, body_font, SIZE_3, False)
    if dt:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.right_indent = Pt(SIZE_3.pt * 2)
        set_line_spacing_28(para)
        para.paragraph_format.first_line_indent = Pt(0)
        run = para.add_run(dt)
        set_cn_font(run, body_font, SIZE_3, False)


ELEMENTS = [
    _render_title,
    _render_salutation,
    _render_body,
    _render_attachments,
    _render_signature,
]